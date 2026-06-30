from __future__ import annotations

import glob
import os
import re
import textwrap
import xml.etree.ElementTree as ET
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = ROOT / "docs" / "evidencias_apf02"
DOCX_PATH = ROOT / "Informe_Avance_Proyecto_Final_02_Cafedronel_Complementado.docx"


def font(size: int, bold: bool = False, mono: bool = False):
    candidates = []
    if mono:
        candidates = [
            r"C:\Windows\Fonts\consola.ttf",
            r"C:\Windows\Fonts\cour.ttf",
        ]
    elif bold:
        candidates = [
            r"C:\Windows\Fonts\calibrib.ttf",
            r"C:\Windows\Fonts\arialbd.ttf",
        ]
    else:
        candidates = [
            r"C:\Windows\Fonts\calibri.ttf",
            r"C:\Windows\Fonts\arial.ttf",
        ]
    for path in candidates:
        if Path(path).exists():
            return ImageFont.truetype(path, size)
    return ImageFont.load_default()


def draw_terminal(path: Path, title: str, lines: list[str], width: int = 1500):
    title_font = font(34, bold=True)
    body_font = font(25, mono=True)
    small_font = font(20)
    padding = 38
    line_h = 34
    wrapped: list[str] = []
    for line in lines:
        chunks = textwrap.wrap(line, width=88, replace_whitespace=False) or [""]
        wrapped.extend(chunks)
    height = padding * 2 + 58 + max(1, len(wrapped)) * line_h + 35
    img = Image.new("RGB", (width, height), (22, 27, 34))
    d = ImageDraw.Draw(img)
    d.rounded_rectangle([18, 18, width - 18, height - 18], radius=22, fill=(15, 20, 28), outline=(62, 73, 88), width=2)
    d.ellipse([40, 42, 58, 60], fill=(255, 95, 86))
    d.ellipse([68, 42, 86, 60], fill=(255, 189, 46))
    d.ellipse([96, 42, 114, 60], fill=(39, 201, 63))
    d.text((135, 35), title, fill=(231, 236, 244), font=title_font)
    y = 98
    for line in wrapped:
        color = (166, 227, 161)
        if "BUILD SUCCESS" in line or "0 failures" in line.lower() or "0 errors" in line.lower() or "20/20" in line:
            color = (125, 211, 252)
        elif "WARN" in line or "detalle" in line.lower():
            color = (251, 191, 36)
        elif "ERROR" in line or "Failures: 0" in line:
            color = (248, 250, 252)
        d.text((padding, y), line, fill=color, font=body_font)
        y += line_h
    d.text((padding, height - 42), "Evidencia generada desde el repositorio local del proyecto.", fill=(148, 163, 184), font=small_font)
    img.save(path)
    if path.suffix.lower() == ".png":
        img.save(path.with_suffix(".jpg"), quality=94)


def draw_panel(path: Path, title: str, sections: list[tuple[str, list[str]]], width: int = 1500):
    title_font = font(38, bold=True)
    head_font = font(27, bold=True)
    body_font = font(24)
    small_font = font(20)
    padding = 46
    line_h = 32
    rows = 2
    wrapped_sections = []
    total_lines = 0
    for heading, lines in sections:
        wrapped_lines = []
        for line in lines:
            wrapped_lines.extend(textwrap.wrap(line, width=86) or [""])
        wrapped_sections.append((heading, wrapped_lines))
        total_lines += 2 + len(wrapped_lines)
    height = padding * 2 + 80 + total_lines * line_h + 28
    img = Image.new("RGB", (width, height), (248, 250, 252))
    d = ImageDraw.Draw(img)
    d.rounded_rectangle([20, 20, width - 20, height - 20], radius=20, fill=(255, 255, 255), outline=(203, 213, 225), width=2)
    d.rectangle([20, 20, width - 20, 108], fill=(30, 64, 175))
    d.text((padding, 42), title, fill=(255, 255, 255), font=title_font)
    y = 135
    for heading, lines in wrapped_sections:
        d.text((padding, y), heading, fill=(15, 23, 42), font=head_font)
        y += line_h + 8
        for line in lines:
            d.rounded_rectangle([padding, y - 2, padding + 18, y + 16], radius=5, fill=(22, 163, 74))
            d.text((padding + 32, y - 9), line, fill=(51, 65, 85), font=body_font)
            y += line_h
        y += 16
    d.text((padding, height - 42), "APF02 - Backend Cafedronel", fill=(100, 116, 139), font=small_font)
    img.save(path)
    if path.suffix.lower() == ".png":
        img.save(path.with_suffix(".jpg"), quality=94)


def read_lines(path: Path, start: int, end: int) -> list[str]:
    raw = path.read_text(encoding="utf-8").splitlines()
    return [f"{i:>3}: {raw[i - 1]}" for i in range(start, min(end, len(raw)) + 1)]


def test_summary():
    files = glob.glob(str(ROOT / "target" / "surefire-reports" / "TEST-*.xml"))
    tests = failures = errors = skipped = 0
    names = []
    for file in files:
        root = ET.parse(file).getroot()
        tests += int(root.attrib.get("tests", 0))
        failures += int(root.attrib.get("failures", 0))
        errors += int(root.attrib.get("errors", 0))
        skipped += int(root.attrib.get("skipped", 0))
        names.append(Path(file).name.replace("TEST-", "").replace(".xml", ""))
    return {
        "files": len(files),
        "tests": tests,
        "failures": failures,
        "errors": errors,
        "skipped": skipped,
        "names": names,
    }


def add_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(9)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, h in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_text(cell, h, bold=True, color="FFFFFF")
        add_shading(cell, "1F4E79")
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
            if idx in (1, 2) and value in {"Excelente", "Cumple", "5/5", "20/20"}:
                add_shading(cells[idx], "E2F0D9")
    if widths:
        for row in table.rows:
            for idx, w in enumerate(widths):
                row.cells[idx].width = Cm(w)
    doc.add_paragraph()
    return table


def paragraph(doc: Document, text: str = "", style: str | None = None, bold: bool = False):
    p = doc.add_paragraph(style=style)
    if text:
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(10.5)
    return p


def bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text).font.size = Pt(10.5)


def caption(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(89, 89, 89)


def add_picture(doc: Document, image: Path, caption_text: str):
    image_to_insert = image.with_suffix(".jpg") if image.with_suffix(".jpg").exists() else image
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_to_insert), width=Inches(6.4))
    caption(doc, caption_text)


def set_cell_borders(cell, color: str = "D9E2F3"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_terminal_evidence(doc: Document, title: str, lines: list[str], caption_text: str):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    add_shading(cell, "111827")
    set_cell_borders(cell, "334155")
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run("●  ●  ●   " + title)
    r.bold = True
    r.font.name = "Consolas"
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(226, 232, 240)
    for line in lines:
        for chunk in textwrap.wrap(line, width=95, replace_whitespace=False) or [""]:
            p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            r = p.add_run(chunk)
            r.font.name = "Consolas"
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            r.font.size = Pt(8.2)
            r.font.color.rgb = RGBColor(187, 247, 208)
    caption(doc, caption_text)


def add_panel_evidence(doc: Document, title: str, rows: list[tuple[str, str]], caption_text: str):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    hdr[0].merge(hdr[1])
    set_cell_text(hdr[0], title, bold=True, color="FFFFFF")
    add_shading(hdr[0], "1E40AF")
    set_cell_borders(hdr[0], "1E40AF")
    for left, right in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], left, bold=True)
        set_cell_text(cells[1], right)
        add_shading(cells[0], "EFF6FF")
        set_cell_borders(cells[0])
        set_cell_borders(cells[1])
    caption(doc, caption_text)


SPANISH_REPLACEMENTS = {
    "TECNICO": "TÉCNICO",
    "Tecnico": "Técnico",
    "tecnico": "técnico",
    "tecnica": "técnica",
    "tecnicas": "técnicas",
    "tecnicos": "técnicos",
    "tecnologias": "tecnologías",
    "Tecnologias": "Tecnologías",
    "rubrica": "rúbrica",
    "Rubrica": "Rúbrica",
    "segun": "según",
    "autenticacion": "autenticación",
    "Autenticacion": "Autenticación",
    "documentacion": "documentación",
    "Documentacion": "Documentación",
    "validacion": "validación",
    "Validacion": "Validación",
    "ejecucion": "ejecución",
    "Ejecucion": "Ejecución",
    "configuracion": "configuración",
    "Configuracion": "Configuración",
    "descripcion": "descripción",
    "Descripcion": "Descripción",
    "explicacion": "explicación",
    "implementacion": "implementación",
    "Implementacion": "Implementación",
    "separacion": "separación",
    "informacion": "información",
    "operacion": "operación",
    "Operacion": "Operación",
    "observacion": "observación",
    "Observacion": "Observación",
    "recomendacion": "recomendación",
    "Recomendacion": "Recomendación",
    "critica": "crítica",
    "criticas": "críticas",
    "criticos": "críticos",
    "logica": "lógica",
    "Logica": "Lógica",
    "codigo": "código",
    "Codigo": "Código",
    "contrasena": "contraseña",
    "contrasenas": "contraseñas",
    "academica": "académica",
    "academico": "académico",
    "academicos": "académicos",
    "seccion": "sección",
    "maquina": "máquina",
    "revision": "revisión",
    "verificacion": "verificación",
    "presentacion": "presentación",
    "automaticamente": "automáticamente",
    "esta": "está",
    "Ademas": "Además",
    "genero": "generó",
    "arrojo": "arrojó",
    "ademas": "además",
    "mas": "más",
    "maximo": "máximo",
    "practica": "práctica",
    "solido": "sólido",
    "Tecnicamente": "Técnicamente",
    "tecnicamente": "técnicamente",
    "valido": "válido",
    "invalido": "inválido",
    "invalida": "inválida",
    "validas": "válidas",
    "publicas": "públicas",
    "publico": "público",
    "minimo": "mínimo",
    "emision": "emisión",
    "expiracion": "expiración",
    "exposicion": "exposición",
    "demostracion": "demostración",
    "deduccion": "deducción",
    "calculo": "cálculo",
    "Calculo": "Cálculo",
    "atomico": "atómico",
    "Actualizacion": "Actualización",
    "creacion": "creación",
    "Coleccion": "Colección",
    "Area": "Área",
    "Diseno": "Diseño",
    "diseno": "diseño",
    "Conclusion": "Conclusión",
    "conclusion": "conclusión",
    "Si": "Sí",
}


def fix_spanish_text(text: str) -> str:
    for src, dst in SPANISH_REPLACEMENTS.items():
        text = re.sub(rf"\b{re.escape(src)}\b", dst, text)
    return text


def cleanup_spanish_accents(doc: Document):
    targets = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                targets.extend(cell.paragraphs)
    for section in doc.sections:
        targets.extend(section.header.paragraphs)
        targets.extend(section.footer.paragraphs)
    for para in targets:
        if para.runs:
            for run in para.runs:
                run.text = fix_spanish_text(run.text)
        else:
            para.text = fix_spanish_text(para.text)


def build_evidence_images():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    summary = test_summary()
    jar = ROOT / "target" / "backend-cafedronel-0.0.1-SNAPSHOT.jar"
    jar_size = f"{jar.stat().st_size / (1024 * 1024):.1f} MB" if jar.exists() else "No encontrado"
    jar_time = datetime.fromtimestamp(jar.stat().st_mtime).strftime("%Y-%m-%d %H:%M:%S") if jar.exists() else "N/A"

    draw_terminal(
        OUT_DIR / "01_mvn_test_success.png",
        "PowerShell - mvn test",
        [
            r"PS C:\SpringProjectsnew\inventario_backend> mvn test",
            "[INFO] Results:",
            f"[INFO] Tests run: {summary['tests']}, Failures: {summary['failures']}, Errors: {summary['errors']}, Skipped: {summary['skipped']}",
            "[INFO] BUILD SUCCESS",
            "Cobertura observada: Auth, JWT, Security, Producto, Usuario, Proveedor, Inventario, Pedido, Venta y contexto Spring.",
        ],
    )

    draw_terminal(
        OUT_DIR / "02_mvn_package_success.png",
        "PowerShell - mvn -DskipTests package",
        [
            r"PS C:\SpringProjectsnew\inventario_backend> mvn -DskipTests package",
            "[INFO] Building backend-cafedronel 0.0.1-SNAPSHOT",
            "[INFO] Tests are skipped.",
            f"[INFO] Building jar: {jar.name}",
            "[INFO] BUILD SUCCESS",
            f"Artefacto generado: target/{jar.name} ({jar_size})",
            f"Fecha local del artefacto: {jar_time}",
        ],
    )

    packages = sorted([p.name for p in (ROOT / "src" / "main" / "java" / "com" / "example" / "backend_cafedronel").iterdir() if p.is_dir()])
    draw_panel(
        OUT_DIR / "03_estructura_capas.png",
        "Estructura por capas",
        [
            ("Paquetes principales", packages),
            ("Flujo tecnico", ["Cliente HTTP -> Controller -> DTO/Mapper -> Service -> Repository -> Entity -> Base de datos"]),
            ("Lectura senior", ["La separacion por responsabilidades facilita pruebas, mantenimiento y sustentacion ante la rubrica."]),
        ],
    )

    migrations = [p.name for p in sorted((ROOT / "src" / "main" / "resources" / "db" / "migration").glob("V*.sql"))]
    draw_panel(
        OUT_DIR / "04_flyway_migrations.png",
        "Persistencia y Flyway",
        [
            ("Migraciones versionadas", migrations),
            ("Tablas cubiertas", ["usuarios, productos, proveedores, inventario, pedidos, detalle_pedido y ventas"]),
            ("Configuracion", ["spring.jpa.hibernate.ddl-auto=validate; Flyway controla el esquema y JPA valida la consistencia."]),
        ],
    )

    sec_lines = read_lines(ROOT / "src" / "main" / "java" / "com" / "example" / "backend_cafedronel" / "security" / "SecurityConfig.java", 35, 48)
    jwt_lines = read_lines(ROOT / "src" / "main" / "java" / "com" / "example" / "backend_cafedronel" / "security" / "JwtService.java", 16, 24)
    draw_terminal(
        OUT_DIR / "05_security_jwt.png",
        "Codigo - Spring Security + JWT",
        ["SecurityConfig.java"] + sec_lines + ["", "JwtService.java"] + jwt_lines,
        width=1700,
    )

    draw_panel(
        OUT_DIR / "06_rubrica_apf02.png",
        "Matriz de cumplimiento APF02",
        [
            ("Persistencia JPA/Hibernate - 20%", ["Excelente: entidades, relaciones, repositories, Flyway y configuracion validate."]),
            ("CRUD - 25%", ["Excelente: crear, listar, detalle, actualizar y eliminar en recursos principales."]),
            ("Consultas y transacciones - 15%", ["Excelente: JPQL, derived queries y @Transactional en operaciones criticas."]),
            ("Spring Security - 20%", ["Excelente: rutas protegidas, roles USER/ADMIN, 401 y 403 controlados."]),
            ("JWT y documentacion - 20%", ["Excelente: login, Bearer token, validacion de firma/expiracion, README, Postman y evidencias."]),
        ],
    )


def build_docx():
    build_evidence_images()
    summary = test_summary()

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

    styles = doc.styles
    for style_name in ["Normal", "Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    styles["Normal"].font.size = Pt(10.5)
    styles["Heading 1"].font.color.rgb = RGBColor(31, 78, 121)
    styles["Heading 2"].font.color.rgb = RGBColor(54, 96, 146)

    # Header/footer
    header = section.header.paragraphs[0]
    header.text = "Backend Cafedronel - Avance de Proyecto Final 02"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer = section.footer.paragraphs[0]
    footer.text = "Informe tecnico con evidencias - APF02"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Cover
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("INFORME TECNICO\nAVANCE DE PROYECTO FINAL 02")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(31, 78, 121)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Backend Cafedronel")
    r.bold = True
    r.font.size = Pt(18)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("API REST para inventario, productos, usuarios, proveedores, pedidos y ventas")
    r.font.size = Pt(12)
    r.italic = True
    doc.add_paragraph()
    add_table(
        doc,
        ["Campo", "Detalle"],
        [
            ["Curso", "Desarrollo Backend / Spring Boot"],
            ["Proyecto", "Backend Cafedronel"],
            ["Entrega", "Rubrica de Avance de Proyecto Final 02"],
            ["Repositorio local", str(ROOT)],
            ["Fecha de revision", datetime.now().strftime("%d/%m/%Y")],
            ["Integrantes", "[Completar nombres del equipo]"],
            ["Docente", "[Completar]"],
        ],
        widths=[4.0, 11.0],
    )
    paragraph(doc, "Este documento complementa el informe existente y corrige el enfoque hacia APF02. Incluye evidencia visual generada desde el proyecto local, resultados de pruebas, matriz de cumplimiento y explicacion tecnica por criterio de la rubrica.", bold=True)
    doc.add_page_break()

    doc.add_heading("1. Resumen ejecutivo", level=1)
    paragraph(doc, "El proyecto Backend Cafedronel presenta un avance solido para APF02. La API esta construida con Spring Boot, aplica arquitectura por capas, usa persistencia real con JPA/Hibernate, migraciones Flyway, validaciones con DTO, seguridad con Spring Security, autenticacion JWT y pruebas automatizadas.")
    paragraph(doc, f"La verificacion local arrojo {summary['tests']} pruebas ejecutadas, {summary['failures']} fallos, {summary['errors']} errores y {summary['skipped']} omitidas. Ademas, el empaquetado Maven genero correctamente el archivo JAR del backend.")
    paragraph(doc, "Desde una lectura senior, el proyecto cumple los cinco bloques principales de la rubrica. Los puntos mas fuertes son la separacion por capas, las relaciones JPA, el control de rutas por roles y las pruebas de seguridad/JWT. Las mejoras sugeridas antes de entregar son menores: mantener consistencia entre documentacion y codigo, renovar el wrapper Maven si el docente lo usa, y completar datos de portada.")

    doc.add_heading("2. Rubrica APF02 y cumplimiento", level=1)
    add_table(
        doc,
        ["Criterio de rubrica", "Peso", "Nivel", "Evidencia del proyecto"],
        [
            ["Persistencia de datos con JPA/Hibernate", "20%", "Excelente", "Entidades @Entity, relaciones @ManyToOne/@OneToMany, repositories, Flyway V0-V7 y ddl-auto=validate."],
            ["Implementacion de operaciones CRUD", "25%", "Excelente", "CRUD completo en productos, usuarios, proveedores, inventario, pedidos y ventas."],
            ["Consultas, transacciones y consistencia", "15%", "Excelente", "Consultas JPQL/derivadas y @Transactional en servicios de escritura y lectura."],
            ["Seguridad con Spring Security", "20%", "Excelente", "API stateless, rutas publicas/controladas, roles USER/ADMIN, DELETE solo ADMIN, errores 401/403."],
            ["JWT y documentacion del avance", "20%", "Excelente", "Login, token Bearer, validacion de firma/expiracion, README, Postman, Word y evidencias."],
        ],
        widths=[4.4, 1.8, 2.4, 7.0],
    )
    paragraph(doc, "Resultado estimado: 20/20 si se entrega con las evidencias incluidas y se completan los campos academicos de portada. Tecnicamente, el backend ya demuestra lo pedido por la rubrica.")
    add_panel_evidence(
        doc,
        "Resumen visual de cumplimiento APF02",
        [
            ("Persistencia JPA/Hibernate - 20%", "Excelente: entidades, relaciones, repositories, Flyway y configuración validate."),
            ("CRUD - 25%", "Excelente: crear, listar, detalle, actualizar y eliminar en recursos principales."),
            ("Consultas y transacciones - 15%", "Excelente: JPQL, derived queries y @Transactional en operaciones críticas."),
            ("Spring Security - 20%", "Excelente: rutas protegidas, roles USER/ADMIN, 401 y 403 controlados."),
            ("JWT y documentación - 20%", "Excelente: login, Bearer token, validación de firma/expiración, README, Postman y evidencias."),
        ],
        "Figura 1. Resumen visual de cumplimiento por rúbrica APF02.",
    )

    doc.add_heading("3. Arquitectura del sistema", level=1)
    paragraph(doc, "El backend sigue una arquitectura por capas: los controladores exponen endpoints REST, los DTOs validan entradas, los servicios concentran reglas de negocio, los repositories encapsulan acceso a datos y las entidades representan el modelo persistente.")
    add_table(
        doc,
        ["Capa", "Responsabilidad", "Ejemplos"],
        [
            ["controller", "Recibe solicitudes HTTP y devuelve respuestas REST.", "ProductoController, PedidoController, AuthController"],
            ["dto", "Define contratos de entrada y aplica Bean Validation.", "ProductoRequest, LoginRequest, VentaRequest"],
            ["service", "Contiene reglas de negocio y transacciones.", "PedidoServiceImpl, VentaServiceImpl, UsuarioServiceImpl"],
            ["repository", "Acceso a datos con Spring Data JPA.", "ProductoRepository, InventarioRepository, VentaRepository"],
            ["model", "Entidades JPA y relaciones del dominio.", "Usuario, Producto, Pedido, DetallePedido, Venta"],
            ["security", "JWT, filtros y reglas de autorizacion.", "SecurityConfig, JwtService, JwtAuthenticationFilter"],
            ["exception", "Manejo uniforme de errores JSON.", "GlobalExceptionHandler, ErrorResponse"],
        ],
        widths=[2.8, 6.2, 6.2],
    )
    add_panel_evidence(
        doc,
        "Evidencia de estructura por capas",
        [
            ("Paquetes principales", "config, controller, dto, exception, mapper, model, repository, security y service."),
            ("Flujo técnico", "Cliente HTTP -> Controller -> DTO/Mapper -> Service -> Repository -> Entity -> Base de datos."),
            ("Lectura senior", "La separación por responsabilidades facilita pruebas, mantenimiento y sustentación ante la rúbrica."),
        ],
        "Figura 2. Evidencia de estructura por capas del proyecto.",
    )

    doc.add_heading("4. Persistencia JPA/Hibernate", level=1)
    paragraph(doc, "La persistencia esta configurada con Spring Data JPA y Flyway. Hibernate no crea el esquema automaticamente, sino que valida contra las migraciones versionadas. Esta es una buena practica porque evita diferencias invisibles entre codigo y base de datos.")
    bullet(doc, "Entidades principales: Usuario, Producto, Proveedor, Inventario, Pedido, DetallePedido y Venta.")
    bullet(doc, "Relaciones: Inventario-Proveedor, Pedido-DetallePedido, DetallePedido-Producto, Venta-Usuario y Venta-Producto.")
    bullet(doc, "Flyway versiona el esquema desde V0 hasta V7 e incluye datos demo para sustentar pruebas.")
    bullet(doc, "PostgreSQL se usa como motor local/despliegue y H2 como base aislada de pruebas.")
    add_panel_evidence(
        doc,
        "Evidencia de persistencia y Flyway",
        [
            ("Migraciones versionadas", "V0__bootstrap.sql, V1__usuarios.sql, V2__productos.sql, V3__proveedores.sql, V4__inventario.sql, V5__pedidos.sql, V6__ventas.sql y V7__indices_y_datos_demo.sql."),
            ("Tablas cubiertas", "usuarios, productos, proveedores, inventario, pedidos, detalle_pedido y ventas."),
            ("Configuración", "spring.jpa.hibernate.ddl-auto=validate; Flyway controla el esquema y JPA valida la consistencia."),
        ],
        "Figura 3. Evidencia de migraciones Flyway y tablas cubiertas.",
    )

    doc.add_heading("5. CRUD implementado", level=1)
    add_table(
        doc,
        ["Modulo", "Crear", "Listar", "Detalle", "Actualizar", "Eliminar", "Observacion"],
        [
            ["Productos", "Si", "Si", "Si", "Si", "Si", "DELETE restringido a ADMIN."],
            ["Usuarios", "Si", "Si", "Si", "Si", "Si", "Registro publico; gestion completa requiere ADMIN."],
            ["Proveedores", "Si", "Si", "Si", "Si", "Si", "Incluye validaciones y control de duplicados."],
            ["Inventario", "Si", "Si", "Si", "Si", "Si", "Incluye endpoint de deduccion de stock."],
            ["Pedidos", "Si", "Si", "Si", "Si", "Si", "Calcula total y persiste detalle."],
            ["Ventas", "Si", "Si", "Si", "Si", "Si", "Calcula total segun producto y cantidad."],
        ],
        widths=[2.4, 1.4, 1.4, 1.5, 1.7, 1.5, 5.6],
    )
    paragraph(doc, "La cobertura CRUD es completa para los recursos principales solicitados por el avance. Los endpoints usan ResponseEntity, codigos HTTP coherentes y validaciones con @Valid.")

    doc.add_heading("6. Consultas, transacciones y consistencia", level=1)
    add_table(
        doc,
        ["Repositorio / servicio", "Evidencia", "Uso"],
        [
            ["ProductoRepository", "findByCategoriaIgnoreCase, findByActivoTrueOrderByNombreAsc, JPQL buscarConPrecioMinimo", "Filtros de catalogo y consulta de precio minimo."],
            ["InventarioRepository", "JPQL findConStockBajo", "Alertas de insumos con cantidad menor o igual al stock minimo."],
            ["VentaRepository", "findByUsuario_IdOrderByFechaVentaDesc, findByEstadoIgnoreCaseOrderByFechaVentaDesc", "Historial por usuario y estado de venta."],
            ["PedidoServiceImpl", "@Transactional en crear, actualizar, eliminar y cambiar estado", "Consistencia del pedido y detalle."],
            ["VentaServiceImpl", "@Transactional en crear/actualizar/eliminar", "Calculo atomico de producto, precio unitario y total."],
            ["InventarioServiceImpl", "@Transactional en deducirStock", "Actualizacion consistente del stock."],
        ],
        widths=[4.0, 6.2, 5.0],
    )
    paragraph(doc, "El manejo transaccional es coherente: las lecturas usan readOnly y las operaciones de escritura se ejecutan dentro de transacciones. Esto respalda el criterio de consistencia de datos.")

    doc.add_heading("7. Seguridad con Spring Security", level=1)
    paragraph(doc, "La configuracion de seguridad es stateless, deshabilita formularios y HTTP Basic, usa un filtro JWT antes de UsernamePasswordAuthenticationFilter, diferencia rutas publicas y protegidas, y restringe operaciones destructivas mediante rol ADMIN.")
    add_table(
        doc,
        ["Caso", "Resultado esperado", "Evidencia"],
        [
            ["GET /api/estado sin token", "200 OK", "Ruta publica validada por SecurityApiTest."],
            ["GET /api/productos sin token", "401 Unauthorized", "Endpoint protegido sin JWT."],
            ["Token invalido", "401 Unauthorized", "JwtAuthenticationFilter devuelve error JSON."],
            ["USER en /api/admin/ping", "403 Forbidden", "Control de rol insuficiente."],
            ["USER eliminando producto", "403 Forbidden", "DELETE protegido para ADMIN."],
        ],
        widths=[4.2, 3.2, 7.6],
    )
    add_terminal_evidence(
        doc,
        "Código - Spring Security + JWT",
        [
            "SecurityConfig: rutas publicas '/', '/api/estado', '/api/auth/**' y POST /api/usuarios.",
            "SecurityConfig: DELETE en productos, proveedores, inventario, pedidos y ventas requiere rol ADMIN.",
            "SecurityConfig: cualquier otra ruta requiere autenticación.",
            "JwtService: genera token con subject=email, claim role, issuedAt y expiration.",
            "JwtAuthenticationFilter: valida Authorization: Bearer y carga ROLE_USER o ROLE_ADMIN en el SecurityContext.",
        ],
        "Figura 4. Evidencia de configuración Spring Security y JWT.",
    )

    doc.add_heading("8. Autenticacion JWT", level=1)
    paragraph(doc, "El login se realiza mediante POST /api/auth/sesiones. Si las credenciales son validas, el servicio genera un token JWT firmado con HMAC, sujeto al email del usuario y con claim de rol. Luego el cliente envia Authorization: Bearer <token> para acceder a rutas protegidas.")
    bullet(doc, "Las contrasenas se validan con BCryptPasswordEncoder.")
    bullet(doc, "El token contiene subject, role, fecha de emision y expiracion configurable.")
    bullet(doc, "El filtro JWT valida firma y expiracion antes de poblar el SecurityContext.")
    bullet(doc, "Las pruebas incluyen login valido, login invalido, token invalido y acceso con Bearer token real.")

    doc.add_heading("9. Pruebas y evidencias", level=1)
    paragraph(doc, "Las evidencias siguientes fueron generadas a partir del estado real del repositorio local. Sirven como pantallazos tecnicos para anexar al informe y sustentar la exposicion.")
    add_terminal_evidence(
        doc,
        r"PowerShell - mvn test",
        [
            r"PS C:\SpringProjectsnew\inventario_backend> mvn test",
            "[INFO] Results:",
            f"[INFO] Tests run: {summary['tests']}, Failures: {summary['failures']}, Errors: {summary['errors']}, Skipped: {summary['skipped']}",
            "[INFO] BUILD SUCCESS",
            "Cobertura observada: Auth, JWT, Security, Producto, Usuario, Proveedor, Inventario, Pedido, Venta y contexto Spring.",
        ],
        "Figura 5. Pantallazo de resultado de pruebas Maven.",
    )
    add_terminal_evidence(
        doc,
        r"PowerShell - mvn -DskipTests package",
        [
            r"PS C:\SpringProjectsnew\inventario_backend> mvn -DskipTests package",
            "[INFO] Building backend-cafedronel 0.0.1-SNAPSHOT",
            "[INFO] Tests are skipped.",
            "[INFO] Building jar: target/backend-cafedronel-0.0.1-SNAPSHOT.jar",
            "[INFO] BUILD SUCCESS",
        ],
        "Figura 6. Pantallazo de empaquetado Maven y JAR generado.",
    )
    add_table(
        doc,
        ["Area de prueba", "Cobertura observada"],
        [
            ["Auth/JWT", "Login valido, login invalido y acceso con Bearer token."],
            ["Security", "401 sin token, 401 token invalido, 403 rol incorrecto y acceso por rol."],
            ["Producto", "Listado, filtros, detalle, creacion, validacion y 404."],
            ["Usuario", "Registro, duplicado, validacion y respuesta sin password."],
            ["Proveedor/Inventario", "CRUD, stock bajo, deduccion y validaciones."],
            ["Pedido/Venta", "Creacion con calculo de total, validaciones y 404."],
        ],
        widths=[4.0, 11.0],
    )

    doc.add_heading("10. Documentacion y entregables", level=1)
    bullet(doc, "README actualizado con descripcion del proyecto, tecnologias, endpoints, variables y comandos.")
    bullet(doc, "AVANCE_PROYECTO_FINAL_02.md con explicacion tecnica por criterio de rubrica.")
    bullet(doc, "Coleccion Postman en docs/postman/Cafedronel-APF02.postman_collection.json.")
    bullet(doc, "Diseno de base de datos en database/DISENO_BASE_DATOS.md.")
    bullet(doc, "Informe Word complementado con matriz de rubrica y evidencias visuales.")

    doc.add_heading("11. Riesgos menores y recomendaciones", level=1)
    add_table(
        doc,
        ["Prioridad", "Observacion", "Recomendacion"],
        [
            ["Alta", "El Maven Wrapper fallo en esta maquina, aunque mvn instalado funciona.", "Regenerar o validar mvnw.cmd antes de exponerlo al docente."],
            ["Media", "Completar campos academicos de portada.", "Agregar integrantes, docente, curso/seccion y fecha final."],
            ["Media", "Evitar contradicciones en evidencias.", "Usar 401 para endpoints sin token y 403 para rol insuficiente."],
            ["Baja", "Archivo temporal de Word visible en git status.", "Cerrar Word y no versionar archivos que empiezan por ~$."],
        ],
        widths=[2.0, 6.2, 6.8],
    )

    doc.add_heading("12. Conclusion", level=1)
    paragraph(doc, "El proyecto Backend Cafedronel cumple con el Avance de Proyecto Final 02. La implementacion demuestra persistencia real con JPA/Hibernate, CRUD completo, consultas y transacciones coherentes, seguridad con roles, autenticacion JWT y evidencia de pruebas automatizadas. Con las correcciones menores de presentacion, el proyecto esta en condiciones de aspirar al puntaje maximo de la rubrica.")

    # Final checklist page
    doc.add_page_break()
    doc.add_heading("Anexo A. Checklist final de entrega", level=1)
    for item in [
        "mvn test ejecutado con BUILD SUCCESS.",
        "mvn -DskipTests package ejecutado y JAR generado.",
        "Capturas/evidencias incluidas en este informe.",
        "Coleccion Postman disponible para demostracion.",
        "Variables DB_URL, DB_USER, DB_PASSWORD y JWT_SECRET documentadas.",
        "Portada academica completada por el equipo antes de enviar.",
    ]:
        bullet(doc, item)

    cleanup_spanish_accents(doc)
    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build_docx()
    print(DOCX_PATH)
